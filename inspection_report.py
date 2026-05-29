#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成督查报告Word文档"""

import random
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def set_cell_border(cell):
    """设置表格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        tcBorder = OxmlElement(f'w:{border_name}')
        tcBorder.set(qn('w:val'), 'single')
        tcBorder.set(qn('w:sz'), '4')
        tcBorder.set(qn('w:color'), '000000')
        tcBorders.append(tcBorder)

    tcPr.append(tcBorders)

def generate_inspection_report():
    """生成督查报告"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # 随机督查信息
    inspection_types = [
        "安全生产督查", "环境保护督查", "工作落实督查",
        "质量管理督查", "食品安全督查", "消防安全督查"
    ]
    inspection_type = random.choice(inspection_types)

    inspection_orgs = [
        "市督查办公室", "省督查组", "专项督查工作组",
        "联合督查小组", "质量监督检验中心"
    ]

    inspected_units = [
        "XX科技有限公司", "XX制造工厂", "XX产业园区",
        "XX建设集团", "XX食品有限公司", "XX化工企业"
    ]

    inspection_org = random.choice(inspection_orgs)
    inspected_unit = random.choice(inspected_units)

    # 随机日期
    base_date = datetime(2026, 5, 1)
    days = random.randint(1, 25)
    inspection_date = base_date.replace(day=days)

    # 督查报告标题
    title = doc.add_heading(f"{inspection_type}报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # 添加基本信息表格
    doc.add_heading("一、基本信息", 1)

    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ["督查机关", inspection_org],
        ["被督查单位", inspected_unit],
        ["督查时间", inspection_date.strftime("%Y年%m月%d日")],
        ["督查类型", inspection_type],
        ["督查组成员", random.choice(["张三组长、李四副组长、王五等3人",
                                       "李华主任、王强副主任等5人",
                                       "赵明队长、孙明等4人"])],
        ["督查编号", f"督字〔2026〕第{random.randint(1,50)}号"]
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        # 第一列加粗
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    # 督查概况
    doc.add_heading("二、督查概况", 1)
    overviews = [
        f"根据年度督查工作计划，{inspection_org}于{inspection_date.strftime('%Y年%m月%d日')}对{inspected_unit}开展了{inspection_type}。督查组通过听取汇报、查阅资料、现场检查、人员访谈等方式，全面了解了该单位在相关方面的工作开展情况。本次督查共检查了{random.randint(8,15)}个具体项目，访谈了{random.randint(5,12)}名相关人员，发现了{random.randint(5,10)}个问题，提出了{random.randint(3,6)}条整改建议。",
        f"按照上级统一部署，{inspection_org}组织督查组于{inspection_date.strftime('%Y年%m月%d日')}赴{inspected_unit}进行{inspection_type}。督查工作严格遵循相关法规和标准要求，采取随机抽查与重点检查相结合的方式。此次督查涵盖了制度建设、日常管理、应急处理等{random.randint(3,6)}个方面，共发现需要注意和改进的地方{random.randint(4,9)}处。",
        f"{inspection_org}依据督查职责要求，于{inspection_date.strftime('%Y年%m月%d日')}对{inspected_unit}组织实施了{inspection_type}。督查过程中坚持实事求是、客观公正的原则，通过多方取证、实地核查等方式深入了解实际情况。本次督查共形成了{random.randint(2,5)}个专项检查记录，涵盖了日常管理、制度建设、人员培训等重要环节。"
    ]
    doc.add_paragraph(random.choice(overviews))

    # 主要成绩
    doc.add_heading("三、主要成绩", 1)
    achievements = [
        [
            "制度建设比较完善。该单位建立健全了相关管理制度和操作规程，明确了各部门和人员的职责分工，为规范管理提供了制度保障。",
            "日常管理较为规范。在日常工作中能够按照相关制度和标准要求执行，各项工作记录基本完整，管理流程相对清晰。",
            "人员培训得到重视。定期组织相关人员参加业务培训和技能学习，提高了员工的专业素质和安全意识。",
            "应急准备工作较为充分。制定了应急预案并配备了必要的应急设备和物资，定期组织应急演练。"
        ],
        [
            "责任体系清晰明确。建立了完善的责任制度，各级责任落实到位，形成了上下联动的工作机制。",
            "工作措施扎实有力。在实际工作中采取了多项有效措施，取得了较为明显的成效，相关工作指标符合要求。",
            "监督检查机制健全。建立了日常检查和定期督查相结合的监督机制，能够及时发现问题并督促整改。",
            "工作作风务实严谨。工作态度认真负责，重视过程管理和细节控制，确保了工作质量。"
        ]
    ]
    selected_achievements = random.choice(achievements)
    for achievement in selected_achievements:
        doc.add_paragraph(achievement, style='List Number')

    # 存在问题
    doc.add_heading("四、存在问题", 1)
    issues_pool = {
        "安全生产督查": [
            "安全培训覆盖不够全面，部分新员工未经过充分的安全教育就上岗操作。",
            "安全隐患排查不够深入，存在部分设备定期检查记录不完整的情况。",
            "应急预案可操作性不强，演练频次不足，员工应急处置能力有待提高。",
            "安全投入不足，部分安全防护设备老化，需要及时更新。",
            "安全责任制落实不到位，个别岗位安全职责不清晰。"
        ],
        "环境保护督查": [
            "环保设施运行记录不够完整，部分时段监测数据缺失。",
            "固废分类收集不够规范，存在混放现象。",
            "环保管理制度执行不够严格，个别环节存在管理漏洞。",
            "环保应急能力不足，应急物资储备不够充分。",
            "环保投入相对不足，减排技术应用推广不够及时。"
        ],
        "工作落实督查": [
            "工作进度不够理想，部分重点任务推进缓慢。",
            "工作质量有待提高，存在重进度轻质量的现象。",
            "协调配合不够有力，部门间沟通协作有待加强。",
            "创新意识不强，工作方法较为传统。",
            "考核激励机制不完善，工作积极性调动不足。"
        ],
        "质量管理督查": [
            "质量意识需要进一步加强，个别员工质量标准掌握不够准确。",
            "质量控制措施执行不够严格，存在操作不规范的情况。",
            "质量问题分析不够深入，改进措施针对性有待提高。",
            "质量记录管理不够规范，追溯体系有待完善。"
        ],
        "食品安全督查": [
            "食品安全管理制度执行不够严格，部分食品采购检验记录不完整。",
            "储存管理不够规范，存在生熟混放、保质期管理不到位的情况。",
            "从业人员健康管理和培训需要进一步加强。",
            "食品安全抽检频次不足，风险防控能力有待提高。"
        ],
        "消防安全督查": [
            "消防设施维护不够及时，部分灭火器过期未按要求更换。",
            "消防通道存在堆放杂物现象，不符合消防安全要求。",
            "消防应急预案操作性不强，演练频次不足。",
            "员工消防意识需要进一步提高，消防知识培训覆盖面不够广泛。"
        ]
    }

    relevant_issues = issues_pool.get(inspection_type, issues_pool["工作落实督查"])
    num_issues = random.randint(3, 5)
    selected_issues = random.sample(relevant_issues, min(num_issues, len(relevant_issues)))

    for issue in selected_issues:
        doc.add_paragraph(issue, style='List Number')

    # 整改建议
    doc.add_heading("五、整改建议", 1)

    suggestions_pool = [
        "严格落实管理责任。各级负责人要切实履行管理职责，加强日常监督检查，确保各项制度规定真正落到实处。",
        "完善制度建设。根据督查发现的问题，进一步健全和完善相关制度，堵塞管理漏洞，形成长效机制。",
        "加强人员培训。制定系统的培训计划，定期组织业务培训和安全教育，提高员工的专业素质和责任意识。",
        "加大投入力度。合理安排资金，及时更新老化设备，改善工作条件，为规范管理和安全生产提供物质保障。",
        "强化监督检查。建立健全监督检查机制，定期开展自查自纠，及时发现和解决问题，防止问题反弹。",
        "建立考核机制。将相关工作完成情况纳入绩效考核，奖优罚劣，充分调动工作积极性。",
        "规范档案管理。完善相关工作记录和档案资料，做到记录真实、完整、规范，便于追溯和检查。",
        "提高应急能力。完善应急预案，加强应急演练，确保发生突发情况时能够及时有效处置。"
    ]

    num_suggestions = random.randint(4, 6)
    selected_suggestions = random.sample(suggestions_pool, num_suggestions)

    for i, suggestion in enumerate(selected_suggestions, 1):
        para = doc.add_paragraph(f"({chr(64+i)}) {suggestion}")

    # 结论
    doc.add_heading("六、督查结论", 1)
    conclusions = [
        f"总体来看，{inspected_unit}在{inspection_type}相关方面做了大量工作，取得了一定成效，但督查过程中发现的问题也需要引起高度重视。建议该单位认真研究分析存在的问题，制定详细的整改方案，明确整改时限和责任人，确保问题整改到位。{inspection_org}将对整改情况进行跟踪督查。",
        f"{inspected_unit}对{inspection_type}工作比较重视，制度建设和管理体系建设取得了一定进展。但督查发现的问题反映出工作中还存在薄弱环节。希望该单位以此次督查为契机，举一反三，全面排查，认真整改，不断提升工作水平。督查组将在规定时间内对整改落实情况开展回头看督查。",
        f"经过督查发现，{inspected_unit}在{inspection_type}方面基本能够按照相关要求开展工作，但与高标准严要求相比仍有一定差距。建议进一步提高思想认识，强化责任担当，采取更加有力的措施推进工作。整改情况要形成书面报告并按时上报。"
    ]
    doc.add_paragraph(random.choice(conclusions))

    # 附件说明
    doc.add_heading("七、附件", 1)
    attachments = random.choice([
        "1. 督查现场照片\n2. 检查记录表\n3. 督查谈话记录\n4. 相关制度文件",
        "1. 督查工作签到表\n2. 现场检查清单\n3. 问题整改清单\n4. 会议记录\n5. 相关资料复印件",
        "1. 督查日程安排\n2. 检查项目清单\n3. 问题清单及证据\n4. 督查组工作记录"
    ])
    doc.add_paragraph(attachments)

    # 签发部分
    doc.add_paragraph()
    sign_para = doc.add_paragraph()
    sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sign_text = f"""督查机关（盖章）：{inspection_org}
督查组组长（签字）：__________
{inspection_date.year}年{inspection_date.month}月{inspection_date.day}日"""

    sign_para.add_run(sign_text)

    # 保存文档
    filename = f"督查报告_{inspected_unit}_{inspection_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    # 简化文件名，去除特殊字符
    filename = filename.replace("XX公司_", "").replace("XX工厂_", "").replace("XX园区_", "").replace("XX集团_", "").replace("XX食品_", "").replace("XX化工_", "")
    filename = filename.replace("XX", "").replace("__", "_")

    save_path = f"F:/cvicse/AIAgent/pi-web/outputs/{filename}"
    doc.save(save_path)
    print(f"✓ 已生成督查报告: {save_path}")
    return filename

if __name__ == "__main__":
    generate_inspection_report()

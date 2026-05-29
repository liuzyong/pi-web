#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成随机内容的Word文档"""

import random
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 随机标题
titles = [
    "人工智能的未来发展趋势",
    "数字化转型的关键成功因素",
    "现代软件架构设计原则",
    "数据驱动的决策制定",
    "云计算技术在企业中的应用"
]

# 随机段落内容
paragraphs_templates = [
    "在当今快速发展的技术环境中，{topic}已经成为企业关注的焦点。研究表明，超过{num}%的领先企业正在积极采用相关技术来提升竞争力。",
    "专家指出，{topic}的发展趋势呈现出明显的加速态势。根据最新的市场调研数据预计，未来五年内该领域的市场复合年增长率将达到{rate}%。",
    "{topic}不仅带来了技术层面的革新，更深刻地影响了商业模式和运营方式。企业需要及时调整战略，以充分利用这一变革带来的机遇。",
    "在实施{topic}的过程中，组织面临着多方面的挑战，包括技术集成、人员培训、成本控制等。成功的案例表明，采用渐进式部署策略可以显著降低风险。",
    "随着技术的不断成熟，{topic}的应用场景正在迅速扩展。从传统行业到新兴领域，各类组织都在探索如何最大化这一技术的价值。",
    "专家建议，企业在推进{topic}时应当建立完善的风险评估和管理机制。同时，加强跨部门协作和知识共享也是确保项目成功的关键因素。",
    "分析显示，{topic}对于提升运营效率和客户满意度具有显著作用。通过实际项目的案例研究，我们可以总结出以下关键经验：首先，明确业务目标；其次，选择合适的技术路径；第三，建立持续优化的机制。",
    "当前，{topic}正处于关键的发展阶段。行业内外的参与者都在积极探索创新应用模式，这使得整个生态圈呈现出蓬勃发展的态势。"
]

topics = ["人工智能", "数字化转型", "云计算", "大数据分析", "机器学习", "区块链技术", "物联网"]

bullet_points_templates = [
    ["明确战略定位和目标", "评估现有技术基础", "制定分阶段实施计划", "建立监测和评估机制"],
    ["优化资源配置", "加强团队建设", "构建技术支撑体系", "完善风险管控"],
    ["提升用户体验", "降低运营成本", "增强数据安全", "促进创新协作"],
    ["建立标准化流程", "加强质量控制", "推动持续改进", "深化知识管理"]
]

def generate_random_document():
    """生成随机Word文档"""
    doc = Document()

    # 设置文档标题
    title = random.choice(titles)
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加日期和作者信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_str = datetime.now().strftime("%Y年%m月%d日")
    info_para.add_run(f"生成日期：{date_str}\n")

    # 添加分隔线
    doc.add_paragraph("=" * 50)

    # 添加简介
    doc.add_heading("一、背景概述", 1)
    intro_template = random.choice(paragraphs_templates)
    intro_text = intro_template.format(
        topic=random.choice(topics),
        num=random.randint(60, 90),
        rate=random.randint(15, 40)
    )
    doc.add_paragraph(intro_text)

    # 添加多个正文段落
    doc.add_heading("二、详细分析", 1)

    num_paragraphs = random.randint(4, 7)
    for i in range(num_paragraphs):
        sub_heading = f"（{i+1}）关键点分析"
        doc.add_heading(sub_heading, 2)

        para_template = random.choice(paragraphs_templates)
        para_text = para_template.format(
            topic=random.choice(topics),
            num=random.randint(40, 85),
            rate=random.randint(10, 45)
        )

        para = doc.add_paragraph(para_text)
        # 随机设置字号
        if random.random() > 0.7:
            para.runs[0].font.size = Pt(random.choice([10, 11, 12]))

    # 添加要点列表
    doc.add_heading("三、核心要点", 1)
    bullets = random.choice(bullet_points_templates)
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # 添加表格
    doc.add_heading("四、数据统计", 1)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'

    # 表头
    headers = ["指标", "当前值", "目标值"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # 表格数据
    for i in range(1, 5):
        table.rows[i].cells[0].text = f"指标 {i}"
        table.rows[i].cells[1].text = str(random.randint(30, 80))
        table.rows[i].cells[2].text = str(random.randint(70, 100))

    # 添加结论
    doc.add_heading("五、结论与建议", 1)
    conclusion_templates = [
        "综上所述，{topic}的发展前景广阔，但同时也面临着诸多挑战。建议企业在推进过程中采取稳健的策略，注重技术积累和人才培养，同时建立完善的评估和优化机制。",
        "通过以上分析可以看出，{topic}已经成为当前发展的重要方向。未来，随着相关技术的不断成熟和应用场景的持续扩展，我们预计将看到更加丰富和深入的应用实践。",
        "{topic}的发展趋势已经十分明显。企业应当抓住这一战略机遇，加快相关布局和能力建设，在激烈的市场竞争中占据有利位置。"
    ]
    conclusion_text = random.choice(conclusion_templates).format(
        topic=random.choice(topics)
    )
    doc.add_paragraph(conclusion_text)

    # 保存文档到 outputs 目录
    filename = f"generated_doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    save_path = f"F:/cvicse/AIAgent/pi-web/outputs/{filename}"
    doc.save(save_path)
    print(f"✓ 已生成文档: {save_path}")
    return filename

if __name__ == "__main__":
    generate_random_document()
